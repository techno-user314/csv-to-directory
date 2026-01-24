from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement, ns

IMAGE_SCALING = 0.8
LEFT_MARGIN = 0.25
RIGHT_MARGIN = 0.25
TOP_MARGIN = 0.25
BOTTOM_MARGIN = 0.5


def _emu_to_twips(emu):
    return int(emu * 1440 / 914400)


def _repeat_table_header(row):
    # Mark row as repeating header
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(ns.qn('w:val'), "true")
    trPr.append(tblHeader)

    # Bold all text in the header row
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.underline = True


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")

    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(ns.qn("w:val"), "nil")
        borders.append(border)

    tblPr.append(borders)


def _shade_row(row, fill_color: str):
    # fill_color: hex color without '#', e.g. 'D9E1F2'
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(ns.qn("w:fill"), fill_color)
        tcPr.append(shd)


def add_preset_page(doc, page_photo_path):
    page = doc.add_paragraph()
    run = page.add_run()
    run.add_picture(page_photo_path, width=Inches(8))


def add_photo_grid_page(doc, photo_paths, photo_labels):
    section = doc.sections[-1]
    section.left_margin = Inches(LEFT_MARGIN)
    section.right_margin = Inches(RIGHT_MARGIN)
    section.top_margin = Inches(TOP_MARGIN)
    section.bottom_margin = Inches(BOTTOM_MARGIN)

    usable_width = (
            section.page_width
            - section.left_margin
            - section.right_margin
    )

    usable_height = (
            section.page_height
            - section.top_margin
            - section.bottom_margin
    )

    cell_width = usable_width / 3
    cell_height = usable_height / 4

    # Square photo size driven by height
    image_size = min(cell_width, cell_height) * IMAGE_SCALING

    table = doc.add_table(rows=4, cols=3)
    table.autofit = False

    photo_index = 0

    for row in table.rows:
        # Set the row height
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(ns.qn('w:val'), str(_emu_to_twips(cell_height)))
        trHeight.set(ns.qn('w:hRule'), 'exact')
        trPr.append(trHeight)

        for cell in row.cells:
            cell.width = cell_width

            if photo_index < len(photo_paths):
                # --- Photo ---
                p_img = cell.paragraphs[0]
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = 0
                p_img.paragraph_format.space_after = 0

                run = p_img.add_run()
                try:
                    run.add_picture(
                        photo_paths[photo_index],
                        width=image_size
                    )
                except:
                    print(photo_paths[photo_index])

                # --- Name labels ---
                p_label = cell.add_paragraph()
                p_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_label.paragraph_format.space_before = 0
                p_label.paragraph_format.space_after = 0

                label = photo_labels[photo_index].split(", ", maxsplit=1)
                lastname_run = p_label.add_run(label[0])
                lastname_run.font.size = Pt(14)
                lastname_run.font.bold = True
                lastname_run.font.underline = True

                rest_run = p_label.add_run(", " + label[1])
                rest_run.font.size = Pt(12)

                photo_index += 1


def add_landscape_table(doc, column_headers, rows):
    section = doc.add_section()
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    section.left_margin = Inches(LEFT_MARGIN)
    section.right_margin = Inches(RIGHT_MARGIN)
    section.top_margin = Inches(TOP_MARGIN)
    section.bottom_margin = Inches(BOTTOM_MARGIN)

    table = doc.add_table(rows=1, cols=len(column_headers))

    available_width = section.page_width - section.left_margin - section.right_margin
    col_width = available_width / len(column_headers)

    header_row = table.rows[0]
    _repeat_table_header(header_row)

    for i, header in enumerate(column_headers):
        cell = header_row.cells[i]
        cell.width = col_width
        run = cell.paragraphs[0].add_run(header)
        run.bold = True

    for idx, row_data in enumerate(rows):
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.width = col_width
            cell.text = text

        if idx % 2 == 0:
            _shade_row(row, "99c1e6")

    _remove_table_borders(table)
